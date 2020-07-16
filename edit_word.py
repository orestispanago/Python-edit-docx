import docx
from docx.shared import Pt

doc = docx.Document("vevaiosi-kykloforias-ergazomenou.docx")


#for count,i in enumerate(range(len(doc.paragraphs))):
#    print(count,doc.paragraphs[i].text)


onoma = "Γεώργιος Τριμπούρδελος"
onoma_patros = "Βρασίδα"
etaireia = "Luben Συνεταιριστική ΕΠΕ"
arithmos_tautotitas = "AAA 123456"
perifereia = "Κάτω Παπαρίτσας"


def add_text(line,text):
    run = doc.paragraphs[line].add_run(text)
    font = run.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.bold = True

add_text(7, onoma)
add_text(8, onoma_patros)
add_text(10, etaireia)
add_text(12, arithmos_tautotitas)
add_text(18, perifereia)

doc.save(f"output/{onoma} του {onoma_patros}.docx")